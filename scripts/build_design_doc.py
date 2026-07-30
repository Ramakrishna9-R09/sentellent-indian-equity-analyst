from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "ARCHITECTURE_DESIGN.md"
OUTPUT = ROOT / "docs" / "Sentellent_Architecture_Design.docx"

NAVY = "17324D"
BLUE = "2E74B5"
MID_BLUE = "DDEBF7"
LIGHT_BLUE = "EAF2F8"
LIGHT_GREY = "F2F4F7"
GREY = "5B6573"
WHITE = "FFFFFF"
BLACK = "000000"
USABLE_WIDTH_DXA = 9360


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths_dxa):
        grid_col.set(qn("w:w"), str(width))

    for row in table.rows:
        row._tr.get_or_add_trPr()
        for cell, width in zip(row.cells, widths_dxa):
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def mark_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, name="Calibri", size=11, color=BLACK, bold=None, italic=None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def set_paragraph_border_bottom(paragraph, color=BLUE) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.76)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11.5, NAVY, 9, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.12

    code = doc.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    code._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    code.font.size = Pt(8.4)
    code.font.color.rgb = RGBColor.from_string(NAVY)
    code.paragraph_format.space_before = Pt(2)
    code.paragraph_format.space_after = Pt(2)
    code.paragraph_format.left_indent = Inches(0.2)
    code.paragraph_format.right_indent = Inches(0.2)


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    p.clear()
    run = p.add_run("SENTELLENT ASSESSMENT  |  ARCHITECTURE DESIGN")
    set_run_font(run, size=8.5, color=GREY, bold=True)
    set_paragraph_border_bottom(p, color="D7DBE2")

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.clear()
    run = p.add_run("Personal Agentic AI Indian Equity Analyst  |  ")
    set_run_font(run, size=8.5, color=GREY)
    add_page_number(p)


def add_inline_markdown(paragraph, text: str, default_size=10.5) -> None:
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=default_size, bold=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=default_size)


def add_cover(doc: Document) -> None:
    doc.add_paragraph().paragraph_format.space_after = Pt(22)
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(6)
    run = kicker.add_run("SENTELLENT HIRING CHALLENGE")
    set_run_font(run, size=11, color=BLUE, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(5)
    title.paragraph_format.line_spacing = 1.02
    run = title.add_run("Personal Agentic AI\nIndian Equity Analyst")
    set_run_font(run, size=26, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    run = subtitle.add_run("Architecture & Delivery Design")
    set_run_font(run, size=15, color=GREY)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(16)
    set_paragraph_border_bottom(rule, color=BLUE)

    lead = doc.add_paragraph()
    lead.paragraph_format.space_after = Pt(18)
    lead.paragraph_format.line_spacing = 1.2
    run = lead.add_run(
        "An implementation-ready plan for a grounded, cited, personalised Indian-equity "
        "research assistant deployed on AWS with Terraform and GitHub Actions."
    )
    set_run_font(run, size=12, color=BLACK)

    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2100, 7260])
    mark_header_row(table.rows[0])
    set_cell_shading(table.rows[0].cells[0], LIGHT_GREY)
    set_cell_shading(table.rows[0].cells[1], LIGHT_GREY)
    for cell, text in zip(table.rows[0].cells, ("Design item", "Commitment")):
        p = cell.paragraphs[0]
        add_inline_markdown(p, text, default_size=9.5)
        for run in p.runs:
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(NAVY)
    metadata = [
        ("Required stack", "Next.js (React), FastAPI (Python), LangChain/LangGraph, PostgreSQL + pgvector, AWS, Terraform, GitHub Actions"),
        ("Primary outcome", "A live AWS RAG vertical slice: login, follow one NSE/BSE ticker, ingest sources, and answer a cited INR question."),
        ("Design principles", "Evidence before claims; profile rules as hard filters; idempotent ingestion; deterministic ranking; automated delivery."),
        ("Assessment readiness", "Includes architecture, data model, graph workflow, AWS plan, CI/CD, tests, delivery phases, and submission proof."),
    ]
    for row, (label, value) in zip(table.rows[1:], metadata):
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        p = row.cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_inline_markdown(p, label, default_size=9.5)
        for run in p.runs:
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(NAVY)
        p = row.cells[1].paragraphs[0]
        add_inline_markdown(p, value, default_size=9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    callout = doc.add_paragraph()
    callout.paragraph_format.space_before = Pt(8)
    callout.paragraph_format.space_after = Pt(0)
    run = callout.add_run("Core promise: no unsupported financial claims, numbers, recommendations, or citations.")
    set_run_font(run, size=10.5, color=NAVY, bold=True)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        cells = [cell.strip() for cell in lines[i].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells):
            rows.append(cells)
        i += 1
    return rows, i


def table_widths(column_count: int) -> list[int]:
    if column_count == 2:
        return [3000, 6360]
    if column_count == 3:
        return [2000, 3360, 4000]
    if column_count == 4:
        return [1500, 2600, 2600, 2660]
    width = USABLE_WIDTH_DXA // column_count
    return [width] * (column_count - 1) + [USABLE_WIDTH_DXA - width * (column_count - 1)]


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    columns = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=columns)
    table.style = "Table Grid"
    set_table_geometry(table, table_widths(columns))
    mark_header_row(table.rows[0])
    for row_index, values in enumerate(rows):
        for col_index in range(columns):
            cell = table.cell(row_index, col_index)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if row_index == 0:
                set_cell_shading(cell, LIGHT_GREY)
                add_inline_markdown(p, values[col_index] if col_index < len(values) else "", default_size=8.7)
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(NAVY)
            else:
                add_inline_markdown(p, values[col_index] if col_index < len(values) else "", default_size=8.4)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def add_code_line(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Code Block")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    set_run_font(run, name="Consolas", size=8.2, color=NAVY)
    set_paragraph_border_bottom(p, color="E1E8F0")


def write_markdown_body(doc: Document) -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    in_code = False
    i = 0
    first_title_seen = False
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()
        stripped = line.strip()

        if stripped.startswith("~~~"):
            in_code = not in_code
            i += 1
            continue
        if in_code:
            add_code_line(doc, line)
            i += 1
            continue
        if stripped.startswith("|"):
            rows, i = parse_table(lines, i)
            add_markdown_table(doc, rows)
            continue
        if not stripped:
            i += 1
            continue
        if stripped.startswith("# "):
            if not first_title_seen:
                first_title_seen = True
            else:
                p = doc.add_paragraph(style="Heading 1")
                add_inline_markdown(p, stripped[2:])
            i += 1
            continue
        if stripped.startswith("## "):
            p = doc.add_paragraph(style="Heading 1")
            add_inline_markdown(p, stripped[3:])
            i += 1
            continue
        if stripped.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            add_inline_markdown(p, stripped[4:])
            i += 1
            continue
        if re.match(r"^\d+\.\s+", stripped):
            p = doc.add_paragraph(style="List Number")
            add_inline_markdown(p, re.sub(r"^\d+\.\s+", "", stripped))
            i += 1
            continue
        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_markdown(p, stripped[2:])
            i += 1
            continue
        p = doc.add_paragraph()
        if stripped.startswith("**") and stripped.endswith("**"):
            add_inline_markdown(p, stripped)
            for run in p.runs:
                run.bold = True
        else:
            add_inline_markdown(p, stripped)
        i += 1


def main() -> None:
    doc = Document()
    style_document(doc)
    add_header_footer(doc)
    add_cover(doc)
    write_markdown_body(doc)
    doc.core_properties.title = "Sentellent Personal Agentic AI Indian Equity Analyst - Architecture & Delivery Design"
    doc.core_properties.subject = "Implementation-ready assessment architecture"
    doc.core_properties.author = "Candidate"
    doc.core_properties.comments = "Generated from the repository design source."
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
